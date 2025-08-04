from datetime import datetime

import docx
import httpx
from fastmcp import FastMCP
from pydantic import Field
import requests
from pypinyin import pinyin, Style

# 初始化MCP服务器
mcp = FastMCP('agent_tool')


# 天气查询处理函数
@mcp.tool(name='weather',description='搜索城市天气')
async def handle_weather(city:str=Field(description='地点的位置信息，例如北京、上海')) -> dict:
    """处理天气查询请求"""

    if not city:
        return {"status": "error", "message": "请提供城市名称"}


    # 模拟天气数据，实际应用中可调用API
    print(f"查询{city}的天气信息")
    return {
        "status": "success",
        "result": f'{city}天气：\n'
                  f'晴，25℃，3级风'
    }


# 数学计算处理函数
@mcp.tool(name='calculator',description='计算具体公式的结果')
def handle_calculation(expression:str=Field(description='需要计算的具体公式')):
    """处理数学计算请求"""

    if not expression:
        return {"status": "error", "message": "请提供计算表达式"}

    try:
        # 安全检查
        if any(op in expression for op in [';', ':', '=', '^', 'import', 'exec', 'eval']):
            return {"status": "error", "message": "不支持该运算表达式"}

        # 替换中文运算符
        expression = expression.replace('加', '+').replace('减', '-')
        expression = expression.replace('乘', '*').replace('除', '/')

        # 使用更安全的计算方式替代eval
        allowed_ops = {'+', '-', '*', '/', '(', ')', '.'}
        for char in expression:
            if not (char.isdigit() or char in allowed_ops):
                return {"status": "error", "message": f"不支持的字符: {char}"}

        # 使用简单的计算逻辑
        result = eval(expression)  # 生产环境中建议使用更安全的计算库
        return {
            "status": "success",
            "result": result
        }
    except Exception as e:
        return {"status": "error", "message": f"计算错误: {str(e)}"}


# Word文档处理函数
@mcp.tool(name='doc writer',description='将输出内容写入文档')
def handle_word(content:str=Field(description='需要写入文档的文本')):
    """处理Word文档操作请求"""

    if not content:
        return {"status": "error", "message": "请提供文档内容"}

    try:
        doc = docx.Document()
        doc.add_heading('由FastMCP Agent创建的文档', 0)
        doc.add_paragraph(content)



        filename = f"agent_mcp.docx"
        doc.save(filename)
        return {
            "status": "success",
            "result": f"文档已保存为 {filename}"
        }
    except Exception as e:
        return {"status": "error", "message": f"Word操作错误: {str(e)}"}


if __name__ == "__main__":
    # Initialize and run the server
    mcp.run(transport='stdio')

